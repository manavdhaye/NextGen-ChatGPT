from langgraph.graph import START,END,StateGraph
from pydantic import BaseModel,Field
from langchain_google_genai import ChatGoogleGenerativeAI
from typing import Literal,TypedDict,List,Annotated,Optional
from langchain_core.messages import SystemMessage, HumanMessage
import operator
from langgraph.types import Send
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from langchain_groq import ChatGroq
import os
import io
from dotenv import load_dotenv
from langchain_tavily import TavilySearch
import re
import time
import comtypes
import comtypes.client
from langchain_core.messages import AIMessage

load_dotenv()

# Initialize your models
model = ChatGroq(
    model="llama-3.3-70b-versatile",
    api_key=os.getenv("GROQ_API_KEY"),
    temperature=0.1, 
)

gemini_model = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash", 
    google_api_key=os.getenv("GOOGLE_API_KEY_DOCX_AGENT")
)
    
class Task(BaseModel):
    id: int
    title: str
    goal: str = Field(..., description="One sentence describing what the reader should understand.")
    subpoints: List[str] = Field(..., min_length=3, max_length=6, description="3 to 6 concrete subpoints to cover.")
    target_words: int = Field(..., description="Target word count (300-500).")
    requires_research: bool = False

class Plan(BaseModel):
    doc_title:str
    audience: str = Field(..., description="Target audience for this document.")
    tone: str = Field(..., description="Writing tone (e.g., professional, academic, technical).")
    output_format: Literal["docx", "pdf", "both","txt"] = Field(..., description="The file format the user requested.")
    tasks:List[Task]

class EvidenceItem(BaseModel):
    title: str
    url: str
    snippet: Optional[str] = None

class EvidencePack(BaseModel):
    evidence: List[EvidenceItem] = Field(default_factory=list)

class RouterDecision(BaseModel):
    needs_research: bool
    queries: List[str] = Field(default_factory=list)

class State(TypedDict):
    topic:str
    saved_filepath:str
    needs_research: bool
    queries: List[str]
    evidence: List[EvidenceItem]
    plan: Optional[Plan]
    sections: Annotated[List[tuple[int, str]], operator.add] 
    merged_md: str
    final_text: str

ORCH_SYSTEM = """You are a Lead Document Architect. Produce a comprehensive outline for a premium, extensive professional report.
Requirements:5-8 sections.
2. Explicitly state in the goals to include data tables, comparisons, and 
1. Adapt the number of sections to the user's request. For a short guide, make 2-4 sections. For a comprehensive report, make deep analysis.
Output STRICTLY matching the Plan schema."""

WORKER_SYSTEM = """You are an Expert Technical Writer. Write ONE extensive, highly detailed section of a professional document.
Constraints:
- Expand on every single bullet point with deep analysis, exa
- You MUST write a massive, comprehensive section. Do NOT summarize.mples, and explanations.
- Include at least ONE Markdown Table (`| Header | Header |`) to organize data beautifully.
- Start with a '## <Section Title>' heading.
- Use **bold** text frequently to emphasize key terms and make the text skimmable.
Output ONLY the section content in clean Markdown."""

ROUTER_SYSTEM = """You are the routing module for an Enterprise Document Generator.
Decide if web research is required to write a highly accurate document on the topic.
- If the topic requires up-to-date facts, statistics, or recent news, set needs_research=True and provide 3-5 highly specific search queries.
- If it is a purely theoretical or historical topic, set needs_research=False."""

RESEARCH_SYSTEM = """You are a Research Synthesizer. Extract the most relevant facts and URLs from the raw search results to create an EvidencePack."""

def router_node(state:State):
    router_decison=gemini_model.with_structured_output(RouterDecision)
    result=router_decison.invoke(
        [
            SystemMessage(content=ROUTER_SYSTEM),
            HumanMessage(content=f"Topic: {state['topic']}")
        ]
    )
    return {"needs_research":result.needs_research,"queries":result.queries}

def route_after_route(state:State):
    if state.get("needs_research", False):
        return "research"
    else :
        return "orchestrator"

def orchestrator(state:State)->dict:
    planner=gemini_model.with_structured_output(Plan)
    evidence=state.get("evidence", [])
    answer=planner.invoke([
        SystemMessage(content=ORCH_SYSTEM),
        HumanMessage(content=f"Topic: {state['topic']}\nEvidence available: {[e.model_dump() for e in evidence]}")
    ])
    return {"plan":answer}

def research(state: State) -> dict:
    print("--- RESEARCHING ---")
    queries = state.get("queries", [])
    tool = TavilySearch(max_results=3)
    
    raw_results = []
    for q in queries:
        raw_results.extend(tool.invoke({"query": q}))
        
    if not raw_results:
        return {"evidence": []}
        
    structured_llm = gemini_model.with_structured_output(EvidencePack)
    processed = structured_llm.invoke([
        SystemMessage(content=RESEARCH_SYSTEM),
        HumanMessage(content=f"Raw Results: {raw_results}")
    ])
    
    # Deduplicate by URL
    dedup = {e.url: e for e in processed.evidence if e.url}
    return {"evidence": list(dedup.values())}

def n_worker_excution(state:State):
    tasklist=[]
    for task in state["plan"].tasks:
        tasklist.append(Send("worker",{"task":task.model_dump(),"topic":state["topic"],"plan":state["plan"].model_dump(),"evidence": [e.model_dump() for e in state.get("evidence", [])]}))
    return tasklist

def worker(payload:dict)->dict:
    topic=payload["topic"]
    plan=Plan(**payload["plan"])
    task=Task(**payload["task"])
    evidence=[EvidenceItem(**e) for e in payload.get("evidence", [])]
    print(f"    -> Writing section: {task.title}")

    doc_title=plan.doc_title
    subpoint_text="\n- " + "\n- ".join(task.subpoints)
    evidence_text = "\n".join([f"- {e.title}: {e.url} ({e.snippet})" for e in evidence]) if evidence else "None"

    section_wise=model.invoke(
        [
            SystemMessage(content=WORKER_SYSTEM),
            HumanMessage(content=f"""
            Document Title: {plan.doc_title}
            Tone: {plan.tone}
            Section: {task.title}
            Goal: {task.goal}
            Subpoints: {subpoint_text}
            Available Evidence: {evidence_text}
            Available Evidence: {evidence_text}
            CRITICAL INSTRUCTION: You MUST write at least {task.target_words} words. 
            You MUST include a Markdown Table in your response to compare data or list features.
            Write in a highly authoritative, premium corporate tone."""
                         )
        ]
    ).content.strip()

    return {"sections": [(task.id, section_wise)]}

def merge_content(state: State) -> dict:
    plan = state["plan"]
    ordered_sections = [md for _, md in sorted(state["sections"], key=lambda x: x[0])]
    body = "\n\n".join(ordered_sections).strip()
    return {"merged_md": f"# {plan.doc_title}\n\n{body}\n"}

def _format_runs(paragraph, text):
    """Helper function to parse **bold** markdown natively into Word"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            # Light Black / Dark Charcoal for highlighted text (softer than pure black)
            run.font.color.rgb = RGBColor(40, 40, 40) 
        else:
            paragraph.add_run(part)

def convert_docx_pdf_(docx_path, pdf_path):
    # Ensure we use absolute paths for Windows COM
    abspath_docx = os.path.abspath(docx_path)
    abspath_pdf = os.path.abspath(pdf_path)
    
    word = None
    try:
        print("--- CONVERTING TO PDF (CLEAN VERSION) ---")
        # Create a Word application object in the background
        comtypes.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        
        # Open the docx
        doc = word.Documents.Open(abspath_docx)
        
        # 17 is the constant for PDF format in Word
        doc.SaveAs(abspath_pdf, FileFormat=17)
        
        doc.Close()
        print(f"✅ SUCCESS! Clean PDF saved at: {pdf_path}")
    except Exception as e:
        print(f"❌ Clean Conversion failed: {e}")
    finally:
        if word:
            word.Quit()
        comtypes.CoUninitialize()

def docx_file(plan:Plan,md:str):
    doc = Document()
    
    # 1. SETUP PAGE MARGINS & SECTIONS
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # 2. SETUP GLOBAL HEADER & FOOTER
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"{plan.doc_title}  |  Confidential & Proprietary"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(10) 
    header_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated by Advanced AI Agent"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(180, 180, 180)

    # 3. SET BASE TYPOGRAPHY (Body Text = 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12) 
    font.color.rgb = RGBColor(70, 70, 70) # Soft dark grey for less eye strain

    # 4. CREATE A BEAUTIFUL COVER PAGE
    doc.add_paragraph("\n\n\n\n") 
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(plan.doc_title)
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(36) 
    title_run.font.bold = True
    title_run.font.underline = True 
    title_run.font.color.rgb = RGBColor(28, 73, 143) 

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"\nTarget Audience: {plan.audience}\nStrategic Tone: {plan.tone.title()}\n\n")
    sub_run.font.size = Pt(14) 
    sub_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break() 

    # 5. PARSE MARKDOWN TO STYLIZED DOCUMENT
    lines = md.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
            
        # Ignore any stray image tags if the AI accidentally hallucinates one
        line = re.sub(r'!\[.*?\]\(.*?\)', '', line).strip()
        if not line:
            i += 1
            continue
            
        # PARSE TABLES 
        if line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 2: 
                cols = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(cols))
                
                table.style = 'Light Shading Accent 1' 
                table.autofit = True
                
                # Fill Headers
                hdr_cells = table.rows[0].cells
                for idx, col_name in enumerate(cols):
                    if idx < len(hdr_cells):
                        p = hdr_cells[idx].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(col_name.replace('**', ''))
                        run.font.bold = True
                        run.font.size = Pt(12) 
                        run.font.color.rgb = RGBColor(40, 40, 40) # Light Black for headers too
                        
                # Fill Data Rows
                for r in range(2, len(table_lines)):
                    row_data = [c.strip() for c in table_lines[r].split('|')[1:-1]]
                    row_cells = table.add_row().cells
                    for idx, val in enumerate(row_data):
                        if idx < len(row_cells):
                            p = row_cells[idx].paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(val.replace('**', ''))
                            run.font.size = Pt(11) 
            continue
            
        # PARSE HEADINGS
        if line.startswith('### '): 
            h = doc.add_heading('', level=3)
            run = h.add_run(line.replace('### ', ''))
            run.font.size = Pt(14) # Subheading = 14pt (Exactly 2pt larger than body text)
            run.font.color.rgb = RGBColor(68, 114, 196) 
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6) 
            
        elif line.startswith('## '): 
            h = doc.add_heading('', level=2)
            run = h.add_run(line.replace('## ', ''))
            run.font.size = Pt(16) # Section Heading = 16pt
            run.font.color.rgb = RGBColor(43, 87, 154) 
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(8) 
            
        elif line.startswith('# '): 
            h = doc.add_heading('', level=1)
            run = h.add_run(line.replace('# ', ''))
            run.font.size = Pt(20) # Main Heading = 20pt
            run.font.color.rgb = RGBColor(31, 56, 100)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12) 
            
        # PARSE BULLET POINTS
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Fully Justified Alignment
            _format_runs(p, line[2:])
                    
        # PARSE REGULAR PARAGRAPHS
        else: 
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(14) 
            p.paragraph_format.line_spacing = 1.25 
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Fully Justified Alignment
            _format_runs(p, line)
                    
        i += 1

    # 6. SAVE (With Timestamp to avoid Permission errors)
    timestamp = int(time.time())
    safe_title = "".join(c if c.isalnum() else "_" for c in plan.doc_title).strip("_").lower()
    filename = f"{safe_title}_{timestamp}.docx"

    export_dir = os.path.join(os.getcwd(), "exports")
    os.makedirs(export_dir, exist_ok=True)
    # This is the path for your Hard Drive
    save_path = os.path.join(export_dir, filename)
    # CRITICAL: Save using 'save_path', not 'filename'!
    doc.save(save_path)
    filepath = f"exports/{filename}"
    print(f"\n✅ SUCCESS! Premium Text Document saved as: {save_path}")
    return f"exports/{filename}"
    


def generate_and_build_pdf(state: State) -> dict:
    print("--- BUILDING PREMIUM GAMMA-STYLE DOCX ---")
    plan = state["plan"]
    md = state["merged_md"]
    docx_filepath=docx_file(plan,md)
    try:
        print("--- CONVERTING TO PDF ---")
        pdf_filepath = docx_filepath.replace(".docx", ".pdf")
        convert_docx_pdf_(docx_filepath, pdf_filepath)
        
        print(f"✅ SUCCESS! PDF version saved as: {pdf_filepath}")
        return {"final_text": md, "saved_filepath": pdf_filepath}
        # pdf_filename = f"{filename}.pdf"
        # convert_docx_pdf_(filename, pdf_filename)
        # print(f"✅ SUCCESS! PDF version saved as: {pdf_filename}")
    except Exception as e:
       print(f"❌ PDF Conversion failed: {e}")
       return {"final_text": f"Error converting to PDF: {e}", "saved_filepath": docx_filepath}

def generate_and_build_docx(state: State) -> dict:
    print("--- BUILDING PREMIUM GAMMA-STYLE DOCX ---")
    plan = state["plan"]
    md = state["merged_md"]
    filepath=docx_file(plan,md)
    return {"final_text": md, "saved_filepath": filepath}

def generate_txt_file(state: State) -> dict:
    print("--- GENERATING PLAIN TEXT FILE ---")
    plan = state["plan"]
    md_content = state["merged_md"]

    # Clean Markdown markers (remove ** and #) for a clean look
    clean_text = md_content.replace("**", "").replace("#", "")
    
    # Create a professional header
    header = f"{'='*60}\n"
    header += f"DOCUMENT: {plan.doc_title.upper()}\n"
    header += f"DATE: {time.ctime()}\n"
    header += f"{'='*60}\n\n"

    # Define filename
    
    timestamp = int(time.time())

    safe_title = "".join(c if c.isalnum() else "_" for c in plan.doc_title).strip("_").lower()
    filename = f"{safe_title}_{timestamp}.txt"
    

    print(filename)
    export_dir = os.path.join(os.getcwd(), "exports")
    os.makedirs(export_dir, exist_ok=True)
    
    # This is the path for your Hard Drive
    save_path = os.path.join(export_dir, filename)
    
    # print("header of text file",header)
    # print("clean_text of text file",clean_text)
    print(len(clean_text))
    print(len(header))

    with open(save_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write(clean_text)
    web_path = f"exports/{filename}"
    return {"final_text": clean_text, "saved_filepath": web_path}

def route_to_builder(state: State):
    choice = state["plan"].output_format
    if "txt" in choice:
        return "generate_txt_file"
    elif "pdf" in choice or "both" in choice:
        return "generate_and_build_pdf"
    elif "docx" in choice:
        return "generate_and_build_docx"
    else:
        return "generate_and_build_docx"


graph=StateGraph(State)
graph.add_node("orchestrator",orchestrator)
graph.add_node("research",research)
graph.add_node("router_node",router_node)
graph.add_node("worker",worker)
graph.add_node("merge_content", merge_content)
graph.add_node("generate_and_build_docx", generate_and_build_docx)
graph.add_node("generate_txt_file",generate_txt_file)
graph.add_node("generate_and_build_pdf",generate_and_build_pdf )


graph.add_edge(START,"router_node")
graph.add_conditional_edges("router_node",route_after_route,{"orchestrator":"orchestrator","research":"research"})
graph.add_edge("research","orchestrator")
graph.add_conditional_edges("orchestrator",n_worker_excution,["worker"])
graph.add_edge("worker", "merge_content")
graph.add_conditional_edges("merge_content",route_to_builder,{"generate_and_build_docx":"generate_and_build_docx","generate_and_build_pdf":"generate_and_build_pdf","generate_txt_file":"generate_txt_file"})
graph.add_edge("generate_and_build_docx", END)
graph.add_edge("generate_and_build_pdf", END)
graph.add_edge("generate_txt_file", END)


# Rename your compiled graph so it doesn't conflict with your main backend workflow
docx_workflow = graph.compile() 
